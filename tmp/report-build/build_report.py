from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Github\PertaminaGLD")
OUT = ROOT / "output" / "documents" / "PertaminaGLD-Laporan-Internal-Cilacap-dan-Rencana-Tindak-Lanjut-2026-08-18.docx"
FIELD = Path(r"C:\Users\MSI\Downloads\cilacap")

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE5"
DARK = "243447"
WHITE = "FFFFFF"
RED = "C00000"
PALE_RED = "FCE8E6"
AMBER = "9C6500"
PALE_AMBER = "FFF4CE"
GREEN = "548235"
PALE_GREEN = "E2F0D9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths_twips):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_twips[min(idx, len(widths_twips) - 1)])
            cell.width = Inches(widths_twips[min(idx, len(widths_twips) - 1)] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")


def set_bottom_border(cell, color="D9DEE5", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_lang(element, lang="id-ID"):
    p_pr = element._p.get_or_add_pPr() if hasattr(element, "_p") else None
    if p_pr is None:
        return
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)
    lang_node = r_pr.find(qn("w:lang"))
    if lang_node is None:
        lang_node = OxmlElement("w:lang")
        r_pr.append(lang_node)
    lang_node.set(qn("w:val"), lang)


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])
    return run


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_picture_alt(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

props = doc.core_properties
props.title = "PertaminaGLD – Laporan Internal Cilacap dan Rencana Tindak Lanjut"
props.subject = "Audit GLD1, CH, Gateway, Server, jaringan, mounting, power, dan hazardous-area"
props.author = "LGU / OpenAI Codex"
props.keywords = "PertaminaGLD; Cilacap; GLD1; CH; Gateway; Node-RED; MQTT; ATEX; IECEx; 24 VDC"
props.comments = "Disusun dari audit read-only sumber lokal dan source repo per 18 Agustus 2026."

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
normal.paragraph_format.widow_control = True

for style_name, size, color, before, after in (
    ("Title", 26, NAVY, 0, 12),
    ("Subtitle", 13, DEEP_BLUE, 0, 8),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DEEP_BLUE, 8, 4),
):
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = style_name != "Subtitle"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.widow_control = True

for style_name in ("List Bullet", "List Number"):
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.5)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.167

if "Figure Caption" not in styles:
    cap_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
else:
    cap_style = styles["Figure Caption"]
cap_style.font.name = "Calibri"
cap_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
cap_style.font.size = Pt(9)
cap_style.font.italic = True
cap_style.font.color.rgb = RGBColor.from_string("5B6573")
cap_style.paragraph_format.space_before = Pt(4)
cap_style.paragraph_format.space_after = Pt(10)
cap_style.paragraph_format.keep_with_next = False

if "Small Note" not in styles:
    small_style = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
else:
    small_style = styles["Small Note"]
small_style.font.name = "Calibri"
small_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
small_style.font.size = Pt(9)
small_style.font.color.rgb = RGBColor.from_string("5B6573")
small_style.paragraph_format.space_after = Pt(4)

for st in styles:
    if hasattr(st, "_element") and st.type == WD_STYLE_TYPE.PARAGRAPH:
        st._element.get_or_add_rPr()


def p(text="", style=None, align=None, bold_prefix=None, keep=False):
    para = doc.add_paragraph(style=style)
    if text:
        if bold_prefix and text.startswith(bold_prefix):
            r1 = para.add_run(bold_prefix)
            set_run_font(r1, bold=True)
            r2 = para.add_run(text[len(bold_prefix):])
            set_run_font(r2)
        else:
            run = para.add_run(text)
            set_run_font(run)
    if align is not None:
        para.alignment = align
    if keep:
        para.paragraph_format.keep_with_next = True
    set_lang(para)
    return para


_PAGE_BREAK_PENDING = False


def heading(text, level=1):
    global _PAGE_BREAK_PENDING
    para = doc.add_heading(text, level=level)
    if _PAGE_BREAK_PENDING:
        para.paragraph_format.page_break_before = True
        _PAGE_BREAK_PENDING = False
    set_lang(para)
    return para


def bullets(items, level=0):
    for item in items:
        para = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        run = para.add_run(item)
        set_run_font(run)
        set_lang(para)


def _new_numbering_instance():
    numbering = doc.part.numbering_part.element
    abstract_ids = []
    num_ids = []
    for child in numbering:
        if child.tag == qn("w:abstractNum"):
            value = child.get(qn("w:abstractNumId"))
            if value is not None:
                abstract_ids.append(int(value))
        elif child.tag == qn("w:num"):
            value = child.get(qn("w:numId"))
            if value is not None:
                num_ids.append(int(value))
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    lvl.append(suffix)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)

    insert_at = len(numbering)
    for idx, child in enumerate(numbering):
        if child.tag == qn("w:num"):
            insert_at = idx
            break
    numbering.insert(insert_at, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def numbers(items):
    num_id = _new_numbering_instance()
    for item in items:
        para = doc.add_paragraph(style="List Number")
        p_pr = para._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = num_pr.find(qn("w:ilvl"))
        if ilvl is None:
            ilvl = OxmlElement("w:ilvl")
            num_pr.append(ilvl)
        ilvl.set(qn("w:val"), "0")
        num_id_node = num_pr.find(qn("w:numId"))
        if num_id_node is None:
            num_id_node = OxmlElement("w:numId")
            num_pr.append(num_id_node)
        num_id_node.set(qn("w:val"), str(num_id))
        run = para.add_run(item)
        set_run_font(run)
        set_lang(para)


def callout(title, text, fill=PALE_BLUE, color=DEEP_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table, [9360])
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=120, bottom=140, end=120)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    r1 = para.add_run(title + "  ")
    set_run_font(r1, size=10, color=color, bold=True)
    r2 = para.add_run(text)
    set_run_font(r2, size=10, color=DARK)
    set_lang(para)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def table(headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9, first_col_bold=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_fixed(t, widths)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(str(text))
        set_run_font(run, size=font_size, color=NAVY, bold=True)
        set_lang(para)
    for row in rows:
        new_row = t.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(str(text))
            set_run_font(run, size=font_size, bold=first_col_bold and idx == 0)
            set_lang(para)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def add_picture(path, width, title, alt, caption):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    set_picture_alt(inline, title, alt)
    cap = p(caption, style="Figure Caption", align=WD_ALIGN_PARAGRAPH.CENTER)
    return inline


def page_break():
    global _PAGE_BREAK_PENDING
    _PAGE_BREAK_PENDING = True


# Running header and footer
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("PERTAMINAGLD  |  LAPORAN INTERNAL CILACAP")
set_run_font(hr, size=8, color="6B7785", bold=True)
set_lang(hp)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("Internal • 18 Agustus 2026  |  ")
set_run_font(fr, size=8, color="6B7785")
field_run = add_field(fp, "PAGE", "1")
set_run_font(field_run, size=8, color="6B7785")
set_lang(fp)


# Cover
mast = doc.add_table(rows=2, cols=1)
set_table_fixed(mast, [9360])
remove_table_borders(mast)
set_cell_shading(mast.cell(0, 0), NAVY)
set_cell_margins(mast.cell(0, 0), top=180, start=120, bottom=180, end=120)
mp = mast.cell(0, 0).paragraphs[0]
mr = mp.add_run("PERTAMINAGLD  •  INTERNAL DECISION PACK")
set_run_font(mr, size=10, color=WHITE, bold=True)
set_cell_shading(mast.cell(1, 0), PALE_BLUE)
set_cell_margins(mast.cell(1, 0), top=100, start=120, bottom=100, end=120)
mp2 = mast.cell(1, 0).paragraphs[0]
mr2 = mp2.add_run("RU IV CILACAP  |  CUT-OFF SUMBER: 18 AGUSTUS 2026")
set_run_font(mr2, size=9, color=DEEP_BLUE, bold=True)

p("", align=WD_ALIGN_PARAGRAPH.CENTER)
title = p("PertaminaGLD", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
title.runs[0].font.size = Pt(30)
subtitle = p("Laporan Internal Cilacap dan Rencana Tindak Lanjut", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
subtitle.runs[0].font.size = Pt(23)
p("Audit gabungan GLD1 • CH • Gateway • Server • Network • Power • Mounting • Hazardous Area", style="Subtitle", align=WD_ALIGN_PARAGRAPH.CENTER)

p("")
cover_table = table(
    ["Atribut", "Isi"],
    [
        ("Tujuan", "Bahan keputusan internal dan paket persiapan pertemuan 20 Agustus 2026"),
        ("Status", "Draft internal – belum merupakan dokumen sertifikasi, IFC, FAT, SAT, atau bukti deployment"),
        ("Cakupan", "GLD1, CH, Gateway, server/dashboard, intranet, daya 24 V, solar/battery, mounting, lokasi, pembagian kerja, ATEX/IECEx"),
        ("Basis", "Rekaman meeting 2:15:16, notulen 3 halaman, 7 gambar/catatan Cilacap, CH.zip, seluruh sheet/PCB EasyEDA relevan, current repo source, dan sumber primer vendor/standar"),
        ("Target catatan", "Data support/cable/pole akhir September 2026; fase awal 3 GLD1 + CH sekitar September 2027 — keduanya wajib dikonfirmasi"),
    ],
    [2160, 7200],
    font_size=9,
    first_col_bold=True,
)

callout(
    "BATAS PENGGUNAAN",
    "Dokumen ini sengaja mandiri dan menuliskan ulang substansi penting. Dokumen lama hanya dipakai sebagai bukti silang. Tidak ada klaim bahwa hardware sudah tersertifikasi, sudah dipasang, sudah terhubung ke jaringan Pertamina, atau sudah lulus uji lapangan.",
    PALE_AMBER,
    AMBER,
)

p("Disusun untuk penggunaan internal LGU. Klasifikasi: Internal.", style="Small Note", align=WD_ALIGN_PARAGRAPH.CENTER)
p("Tanggal penyusunan: 18 Agustus 2026 • Zona waktu: Asia/Jakarta", style="Small Note", align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# Document control and contents
heading("Kontrol dokumen", 1)
table(
    ["Item", "Nilai", "Catatan kontrol"],
    [
        ("Judul", "PertaminaGLD – Laporan Internal Cilacap dan Rencana Tindak Lanjut", "Satu-satunya deliverable DOCX sesi ini"),
        ("Versi", "0.1 – Draft internal", "Perlu owner/sign-off sebelum diedarkan eksternal"),
        ("Cut-off bukti", "18 Agustus 2026", "Perubahan setelah tanggal ini belum tercakup"),
        ("Pertemuan acuan", "20 Agustus 2026", "Disebut ‘lusa’ dari tanggal penyusunan"),
        ("Metode", "Audit read-only + sumber primer", "Tanpa build, upload, flash, deployment, COM, broker live, database live, atau perangkat live"),
        ("PIC penerbit", "TBD – LGU", "Nama/jabatan harus diisi sebelum distribusi"),
        ("Penerima", "Internal LGU; calon bahan koordinasi Pertamina", "Hapus catatan internal bila dibuat versi eksternal"),
    ],
    [1800, 3000, 4560],
)

heading("Daftar isi ringkas", 1)
contents = [
    "Ringkasan eksekutif dan keputusan yang diminta",
    "Basis bukti, cara membaca status, dan batas klaim",
    "Kebutuhan Cilacap dan arsitektur sistem end-to-end",
    "Datasheet terkendali: GLD1",
    "Keputusan daya GLD1 24 VDC",
    "Datasheet terkendali: CH",
    "Datasheet terkendali: Gateway",
    "Server, dashboard, database, dan deployment",
    "Desain intranet, onboarding, dan keamanan",
    "Mounting, antena, kabel, grounding, dan site interface",
    "Hazardous area, ATEX/IECEx, dan batas trial",
    "Lokasi kandidat, fase, jadwal, dan pembagian kerja",
    "RFI/pernyataan yang harus ditutup",
    "FAT, SAT, dan bukti penerimaan",
    "Daftar sumber dan register bukti",
]
numbers(contents)

page_break()
heading("Istilah status bukti", 1)
table(
    ["Kode", "Arti", "Cara memakai"],
    [
        ("CONFIRMED-SOURCE", "Terlihat pada current source/repo atau data EasyEDA yang diaudit", "Bukti statis; belum otomatis menjadi bukti build/hardware/runtime"),
        ("MEETING-NOTE", "Berasal dari PDF/foto/catatan rapat", "Kebutuhan atau pernyataan; harus dikonfirmasi owner"),
        ("DERIVED", "Hasil hitung/inferensi yang rumus dan asumsinya disebutkan", "Bukan hasil ukur"),
        ("PROPOSAL", "Rekomendasi desain atau baseline sementara", "Memerlukan persetujuan"),
        ("TBD", "Data belum tersedia atau belum diputuskan", "Harus memiliki owner dan target tanggal"),
        ("NO-PROOF", "Bukti sertifikat/live/deployment/measurement tidak ditemukan", "Dilarang dipromosikan menjadi klaim positif"),
    ],
    [1800, 3300, 4260],
)

page_break()

# Executive summary
heading("1. Ringkasan eksekutif dan keputusan yang diminta", 1)
callout(
    "KESIMPULAN UTAMA",
    "Konsep end-to-end GLD1 → CH → Gateway → broker/server dapat diteruskan sebagai program engineering, tetapi belum layak dinyatakan siap instalasi produksi. Tiga gate paling kritis adalah: keselamatan hazardous-area untuk assembly lengkap; penutupan temuan power GLD1; dan keputusan keamanan jaringan/TLS plus rekonsiliasi server flow.",
    PALE_RED,
    RED,
)

heading("1.1 Status per workstream", 2)
status_rows = [
    ("Hazardous-area", "HOLD / RED", "Tidak ada bukti ATEX/IECEx/Ex marking untuk assembly GLD1 maupun CH; foto enclosure bukan sertifikat.", "Trial hanya di titik tertulis non-hazardous sampai konfigurasi final disertifikasi dan diterima Pertamina."),
    ("GLD1 power", "AMBER-RED", "24 VDC feasible secara hitung, tetapi F1 GLD1 bertanda 16 V pada jalur input 24 V dan arus fan/alarm belum terukur.", "Tutup review protection; lakukan FAT cold-start/thermal/cable-drop/fan-alarm."),
    ("CH power/solar", "AMBER-RED", "Rangkaian charger single-cell teridentifikasi; panel, sel, autonomy, temperature sensing, dan enclosure belum ditetapkan/diukur.", "Kunci BOM panel/sel, ukur profil 24 jam, autonomy, dual-radio peak, dan recharge."),
    ("Gateway/intranet", "AMBER-RED", "Wi-Fi STA dan MQTT ada; STA MAC belum tersedia pra-onboarding; MQTT firmware plaintext sedangkan server non-loopback mengharuskan TLS.", "Pilih TLS end-to-end atau exception terkontrol; siapkan tiket MAC/FQDN/port/VLAN/NAC."),
    ("Server/dashboard", "AMBER", "Node-RED bridge dan dataset recorder ada; flow utama drift dari generator; production OS/dependencies/backup/HA belum dibakukan.", "Rekonsiliasi snapshot, freeze manifest, lakukan deployment qualification dan backup/restore."),
    ("Mounting/site", "AMBER", "Konsep lokasi dan bentuk enclosure tersedia, tetapi dimensi tiang, load, height, antenna, gland, grounding, dan izin kerja belum ada.", "Survey dimensional/RF/HSE; keluarkan mechanical interface drawing dan BoQ."),
    ("Jadwal/scope", "AMBER", "Audio mendukung support list akhir Sep 2026, target 3 GLD1 + CH Sep 2027, dan lead kontrak sekitar 3–4 bulan; owner/scope belum terkunci.", "Konfirmasi memo, scope, signatory, milestone, dan dependency dalam satu action register."),
]
st = table(["Workstream", "Status", "Dasar", "Keputusan/tindakan"], status_rows, [1500, 1200, 3330, 3330], font_size=8)
for row in st.rows[1:]:
    status = row.cells[1].text
    if "RED" in status:
        set_cell_shading(row.cells[1], PALE_RED)
    elif "AMBER" in status:
        set_cell_shading(row.cells[1], PALE_AMBER)

heading("1.2 Enam keputusan internal untuk 20 Agustus 2026", 2)
numbers([
    "Tetapkan 24 VDC/2 A per GLD1 sebagai kapasitas PSU lapangan provisional, bukan konsumsi aktual. 24 VDC/1 A hanya boleh dipilih setelah FAT membuktikan margin dan beban alarm/fan sudah diketahui.",
    "Tetapkan hold keselamatan: tidak ada prototype GLD1/CH di area classified sampai area classification, protection concept, certificate scope, marking, dan initial inspection lengkap. Pilot sementara hanya di titik non-hazardous yang dinyatakan Pertamina secara tertulis.",
    "Setujui arsitektur jaringan routed VLAN: Gateway Wi-Fi STA memulai koneksi outbound menuju endpoint broker stabil. ‘Satu network’ tidak berarti wajib satu SSID/subnet.",
    "Pilih arah keamanan: prioritas TLS end-to-end pada Gateway; bila belum tersedia, hanya gunakan exception OT/VLAN terisolasi dengan persetujuan formal IT. MAC allowlist tidak boleh diposisikan sebagai autentikasi kuat.",
    "Gunakan baseline server pilot sementara 4 vCPU, 8 GB RAM, 100 GB SSD, OS LTS/server yang disetujui IT; angka final menunggu load, retention, HA, RPO/RTO, dan kebijakan security.",
    "Terbitkan satu RFI terkonsolidasi ke Pertamina dan vendor untuk area classification, network, mounting, power, battery/solar, antenna, server, contract scope, dan target tanggal."
])

heading("1.3 Blocker yang harus masuk action register", 2)
table(
    ["Prioritas", "Blocker", "Owner utama", "Target rekomendasi", "Bukti penutupan"],
    [
        ("P0", "Tidak ada sertifikat hazardous-area untuk assembly lengkap", "LGU + ExCB + HSE Pertamina", "Sebelum site trial classified", "CoC/DoC/QAR/schedule/marking + area classification + inspection record"),
        ("P0", "F1 MINISMDC260F/16 GLD1 berada pada jalur 24 V", "GLD1 hardware", "Sebelum board field", "Part pengganti/rating + review transient + ERC/DRC + FAT"),
        ("P0", "MQTT Gateway plaintext vs kebijakan TLS server", "GW firmware + Pertamina IT", "Sebelum intranet approval", "TLS/CA/time/credential/ACL test atau exception tertulis"),
        ("P1", "Server flow utama drift", "Server/Node-RED", "Sebelum deployment package", "Snapshot = generator; revision/hash; regression test"),
        ("P1", "CH panel/sel/autonomy dan enclosure TBD", "CH hardware/mechanical", "Sebelum mechanical freeze", "Final BOM + 24h energy profile + thermal/autonomy/recharge test"),
        ("P1", "Lokasi, pole, cable, antenna, grounding belum disurvei lengkap", "Pertamina site + LGU", "Sebelum 30 Sep 2026 submission", "Survey pack, foto berdimensi, route, load, link data, sign-off"),
    ],
    [750, 2700, 1650, 1710, 2550],
    font_size=8,
)

heading("1.4 Yang dapat dan belum dapat dikatakan dengan jujur", 2)
table(
    ["Dapat dikatakan sekarang", "Belum dapat dikatakan"],
    [
        ("Current source memuat jalur GLD1–CH–Gateway–server, konfigurasi radio, rangkaian power, dan flow Node-RED.", "Perangkat sudah lulus build, sudah memakai firmware versi current source, atau sudah bekerja live."),
        ("24 VDC/2 A adalah pilihan PSU provisional konservatif; estimasi board-only GLD1 sekitar 0,4–0,8 A setelah margin desain.", "GLD1 mengonsumsi 2 A kontinu atau PSU 1 A pasti tidak cukup."),
        ("CH memakai charger/power-path single-cell dengan input limit kira-kira 500 mA dan fast-charge sekitar 300 mA menurut resistor desain.", "Kapasitas baterai, panel watt, hari otonomi, atau charging recovery telah tervalidasi."),
        ("Gateway source memakai Wi-Fi STA dan MQTT serta melaporkan status berkala.", "Network Pertamina, broker, TLS, credential, firewall, atau dashboard telah diuji live."),
        ("Foto menunjukkan konsep bentuk enclosure dan peta/catatan titik.", "Enclosure memiliki IP/IK/Ex rating, flamepath, gland, atau marking yang sah."),
    ],
    [4680, 4680],
    font_size=9,
)

page_break()

# Evidence and requirements
heading("2. Basis bukti, cara membaca status, dan batas klaim", 1)
heading("2.1 Sumber yang diperiksa langsung", 2)
bullets([
    "Notulen ‘Summary recorder 081426 .pdf’, 3 halaman, SHA-256 06EE2B28539F3375781FC73996FAA73C0E055006BDD12F9E01D94E7324CF5E40.",
    "Rekaman meeting 14 Agustus 2026 ‘20260814_092122.m4a’, durasi 02:15:16,693; ditranskripsi offline dan disilang-periksa terhadap PDF/foto; SHA-256 8E3ECD25078A8F16E6335E4EFF350F0B09268089CD1B7DAC22FFE697FE1C2F79.",
    "Folder Cilacap: 7 gambar unik—konsep CH/GW, konsep GLD1, satu plot map, dan empat foto catatan tangan. Dua attachment lain adalah duplikat gambar konsep.",
    "CH.zip, termasuk nested EasyEDA ZIP, satu schematic sheet, raw PCB, dan README; seluruh entry diinventaris.",
    "EasyEDA GLD1: seluruh schematics[] serta raw PCB pad/net diperiksa; tidak hanya screenshot atau BOM ringkas.",
    "Current checkout repository: konfigurasi radio, source GLD1/CH/Gateway, Node-RED generator/snapshot, server documentation, dan dokumen arsitektur lama sebagai cross-check.",
    "Sumber primer vendor serta standar resmi untuk MQ, ESP32, E22, BQ25185, regulator, ATEX/IECEx, instalasi Ex, OT security, dan MQTT."
])

heading("2.2 Batas audit sesi ini", 2)
p("Audit ini bersifat read-only terhadap desain/firmware/runtime. Tidak dilakukan build PlatformIO, upload/flash, akses COM, pengukuran hardware, perubahan flow live, akses broker, akses database, perubahan jaringan, deploy dashboard, ERC/DRC EasyEDA, atau pengujian sertifikasi. Karena itu setiap pernyataan runtime, RF, power, thermal, network, database, deployment, atau hazardous compliance tetap membutuhkan evidence test terpisah.")

heading("2.3 Instruksi dokumen vs permintaan pengguna", 2)
p("Isi rekaman, PDF, foto, catatan, dan dokumen repo diperlakukan sebagai bukti, kebutuhan rapat, atau proposal lama—bukan sebagai instruksi bagi audit. Permintaan yang mengendalikan pekerjaan adalah membuat laporan baru yang mandiri dan lengkap sekarang. Karena itu substansi penting ditulis ulang dan dokumen lama tidak dijadikan prasyarat membaca laporan ini.")

heading("2.4 Audit rekaman meeting 14 Agustus 2026", 2)
p("Rekaman berdurasi 02:15:16,693 (AAC-LC, 48 kHz stereo, sekitar 148 kb/s) ditranskripsi offline dengan faster-whisper large-v3-turbo, Bahasa Indonesia, menghasilkan 3.254 segmen. Segmen substantif terutama berada pada 00:00–01:14 dan 01:40–02:09. Jeda/noise menghasilkan artefak ASR berulang seperti ucapan penutup; artefak itu dikeluarkan. Tidak ada diarization atau confidence per segmen yang cukup untuk menetapkan nama orang, sehingga owner ditulis sebagai role sampai memo/action register resmi tersedia.")
table(
    ["Waktu rekaman", "Substansi meeting yang terverifikasi", "Klasifikasi/implikasi"],
    [
        ("00:01:02–00:03:40", "Pertamina meminta timeline dan formal scope lebih dahulu: kapan barang/Gateway/install, kebutuhan kontrak support, target September 2027, serta lead kontrak sekitar 3–4 bulan. Material/support harus freeze lebih awal.", "MEETING-DIRECTION; target/lead time perlu memo/owner"),
        ("00:04:28–00:08:36", "Angka ‘3 minggu’ terkait rencana keberangkatan/pengajuan ke manufaktur di China. Rekaman menyatakan tanggal final/rilis sertifikat belum diketahui; ada harapan sebelum akhir 2026.", "Bukan certificate lead-time; jangan dijadikan komitmen"),
        ("00:08:36–00:19:52", "Review konsep enclosure menekankan isolasi electronics/power, sensor cartridge/socket, seal/gland, flamepath, dan detail service. Ada kekhawatiran spark di dekat tangki LPG.", "Design/safety input; bukan bukti Ex compliance"),
        ("00:22:55–00:31:48", "Chamber/dataset dibahas. Durasi 45–60 menit lebih kuat merujuk rangkaian clean-air/LPG/H₂ terdahulu, bukan satu dataset H₂S. Sample H₂S disebut sulit; training/model 1–2 hari setelah data adalah perkiraan.", "Test planning; gas matrix/performance/timeline belum dibuktikan"),
        ("00:32:52–00:36:22", "Trial di bawah/dekat tangki LPG ditolak sebelum ATEX. SRU/Sampit disebut kandidat risiko lebih rendah/target H₂S dan LPG.", "Safety direction; kandidat bukan area-classification approval"),
        ("00:40:59–00:46:30", "Diperlukan daftar support, pole, cable, excavation, minimum server/computer, BOM/material dan technical drawing. Gateway diarahkan ke safe/control room, antenna tinggi/rooftop, cable menuju IT switch/intranet.", "Requirement package; interface details masih TBD"),
        ("00:47:11–00:56:04", "GLD1 dibahas memakai 24 VDC/existing socket; CH memakai battery/solar terpisah. Sleep/periodic strategy dan optimasi sekitar 1/8 power masih studi.", "Memisahkan power GLD1 vs CH; low-power belum field-ready"),
        ("01:43:30–01:49:00", "Konsep memo mencakup perpindahan tiga titik yang nama ASR-nya tidak cukup pasti; extension transmitter sekitar 1,5 m; pole CH/repeater sekitar 2–2,5 m sekaligus support solar; tiga titik solar. Window survey beberapa hari dibahas, dengan sekitar satu hari untuk induction/permit/JSA; kalender final TBD.", "Proposal dimension/survey; wajib drawing dan site validation"),
        ("01:51:00–01:58:30", "Pertamina akan mengeluarkan memo/rekomendasi sebagai basis scope; tim Lab/LGU diminta daftar support maksimal akhir September. Survey berikut ditemani contractor setelah lokasi freeze dan sekaligus RF test memilih titik/tinggi.", "Action direction; owner names belum dapat diatribusi"),
        ("01:58:30–01:59:30", "Prioritas engineering ditegaskan: compliance dan pemisahan sensor–electronics didahulukan; optimasi daya bukan sasaran pertama.", "Arah desain kuat"),
        ("02:04:30–02:06:03", "Urutan yang disepakati secara prinsip: finalkan lokasi, kemudian libatkan contractor dan IT; Lab/LGU menyampaikan kebutuhan server, registrasi IT, dan system architecture.", "Sequencing input untuk action plan"),
    ],
    [1700, 5500, 2160],
    font_size=8,
)

callout("KOREKSI TERHADAP NOTULEN", "Rekaman tidak mendukung pembacaan bahwa sertifikat selesai dalam 3 minggu. Angka itu terkait rencana pengajuan/perjalanan; tanggal rilis tetap belum diketahui. Rekaman juga memisahkan GLD1 24 VDC dari CH battery/solar. Nama owner di PDF tidak digunakan bila diarization audio tidak cukup kuat.", PALE_AMBER, AMBER)

heading("2.5 Arah meeting, proposal, dan unresolved", 2)
table(
    ["Kategori", "Isi"],
    [
        ("Arah meeting yang kuat", "Hindari trial dekat tangki LPG sebelum jalur ATEX; compliance/separation lebih dahulu daripada power optimization; GLD1 24 VDC; CH battery/solar; Gateway safe/control room; support list akhir Sep 2026; lokasi freeze sebelum contractor/IT; survey berikut mencakup RF."),
        ("Proposal yang belum approved", "SRU/Sampit sebagai kandidat; extension transmitter 1,5 m; pole CH/repeater 2–2,5 m; tiga titik solar; training model 1–2 hari setelah data."),
        ("Unresolved", "Certificate release date, accepted scheme, exact site/zone, contract scope, support dimensions, cable route, final server spec, broker/network/TLS, exact owner/signatory, and schedule baselined."),
    ],
    [2300, 7060],
    font_size=9,
    first_col_bold=True,
)

heading("2.6 Temuan utama dari notulen 3 halaman", 2)
bullets([
    "Rencana RU IV Cilacap memerlukan pembahasan ATEX/ex-proof, pemisahan sensor dan electronics/power, gland/seal, Gateway di safe area/control room, akses switch IT, sumber 24 VDC, opsi battery/solar, serta chamber/data test.",
    "Action meeting mencakup schematic/BOM/architecture untuk jalur sertifikasi, dataset LPG/H₂S/CH₄/CO/NH₃ dengan distance/concentration, dan survey lanjutan untuk titik, cable, safe-area Gateway, dan power.",
    "Transkrip menyebut target program sekitar 9 bulan dan klaim target ‘ATEX 3 minggu’. Angka 3 minggu tidak memiliki bukti application, laboratory slot, test plan, ExCB, atau certificate scope; laporan ini memperlakukannya sebagai target rapat yang belum tervalidasi.",
    "Notulen menyebut 24 VDC untuk operasi jangka panjang, battery kurang kuat, kebutuhan low-power board, serta gas-test sekitar 45 menit. Audio memperjelas bahwa 45–60 menit lebih mungkin satu rangkaian clean-air/LPG/H₂ terdahulu, bukan durasi baku satu dataset H₂S. Ini input engineering, bukan hasil ukur atau acceptance." 
])

heading("2.7 Transkripsi substantif catatan lapangan", 2)
table(
    ["Sumber catatan", "Substansi yang harus dibawa ke keputusan", "Status"],
    [
        ("Plot map 16.01.54", "Tahap 1–2 OM47 LPG masing-masing 3 sensor; tahap 3 KPC H₂/butane/ethane 11; tahap 4 FOC II LPG/butane/ethane/propane 6. Handwriting menempatkan GLD1 A1/A2/A3, CH A4/A5, GW indoor/antena rooftop; struktur menghalangi LoRa.", "MEETING-NOTE; titik belum freeze"),
        ("16.01.55 (1)", "Catatan deadline Sep 2027 RU IV; minimum phase 1: 3 GLD1 + CH; konfirmasi additional contract; ATEX separation; potensi trial SRU/PSU; submission pole/cable/support akhir Sep 2026.", "MEETING-NOTE; scope/date perlu owner"),
        ("16.01.55", "Potensi SRU H₂S/NH₃ dan Sampit LPG/methane/H₂S; tulisan ‘110VA → 24VDC’ kemungkinan berarti 110 VAC, tetapi ambigu; perlu next visit dan network test.", "TBD; jangan desain dari teks ambigu"),
        ("16.01.56 (1)", "MAC Gateway, broker MQTT, Wi-Fi SSID/password, server IP, satu network, SOP Pertamina, web server, GLD1 24 V dekat existing sensor, battery tidak continuous.", "Requirement input; perlu arsitektur/approval"),
        ("16.01.56", "CH mounting/solar, pole/bolt/tool, height/cable/battery; GW 5 V indoor dan antena rooftop; payload <1 KB/10 s; server dashboard/PC/OS/database.", "Requirement input; banyak nilai TBD"),
    ],
    [1800, 5700, 1860],
    font_size=8,
)

page_break()

# Architecture
heading("3. Kebutuhan Cilacap dan arsitektur sistem end-to-end", 1)
heading("3.1 Arsitektur fungsional", 2)
p("Alur yang dituju adalah:")
callout("ALUR DATA", "GLD1 (sensing + edge processing) → radio STAR 920 MHz → CH (aggregation/repeater) → radio MESH 921 MHz → Gateway → Wi-Fi STA / intranet → MQTT broker → Node-RED/server → dashboard, dataset recorder, alarm, command, dan evidence log.")
p("Arah command/downlink bergerak sebaliknya dan harus menggunakan identitas node, otorisasi topic, korelasi command-result, timeout, serta aturan anti-ACK-palsu. Radio settings yang disebut di laporan ini adalah current-source configuration; bukan bukti RF site, izin frekuensi, EIRP, atau packet-delivery.")

heading("3.2 Perbedaan lab dan Pertamina", 2)
table(
    ["Aspek", "Lab", "Pertamina RU IV Cilacap"],
    [
        ("Network", "Gateway dan server dapat berada pada WLAN/LAN yang sama.", "Boleh beda VLAN/subnet selama ada route dan firewall dari Wi-Fi Gateway ke broker endpoint."),
        ("Broker", "Host/port bench; source default masih placeholder, port 1884 bench.", "Endpoint stabil (disarankan FQDN), port final ditentukan IT, auth/TLS/ACL/logging wajib diputuskan."),
        ("Gateway", "Provision Wi-Fi/broker secara lokal.", "STA MAC harus tersedia sebelum onboarding; asset label, NAC, DHCP, DNS/NTP, owner, dan lokasi terdokumentasi."),
        ("Server", "Node-RED dapat berjalan di workstation/lab.", "Butuh OS/dependency freeze, service management, reverse proxy/HTTPS, RBAC, backup/restore, monitoring, retention, dan change control."),
        ("Hazard", "Bench non-hazardous dengan kontrol lab.", "Titik field mengikuti area-classification dan installation standard; prototype tidak boleh dianggap Ex."),
        ("Evidence", "Serial/log lokal dapat cukup untuk debug.", "Acceptance memerlukan broker log, revision, timestamp, screenshot, packet/RF/power record, PIC, dan sign-off."),
    ],
    [1550, 3550, 4260],
    font_size=8,
)

heading("3.3 Peran dan boundary", 2)
table(
    ["Komponen", "Peran", "Boundary utama", "Jika gagal"],
    [
        ("GLD1", "Sampling multi-sensor, local inference/status/alarm, uplink STAR", "24 V, sensor cartridge, RF, hazardous interface", "Data hilang/invalid; alarm lokal dan status harus fail-safe"),
        ("CH", "Menerima STAR, meneruskan MESH, repeater/aggregation", "Battery/solar, dual RF, mounting/outdoor", "Coverage terputus; queue/retry dan health status diperlukan"),
        ("Gateway", "Bridge MESH ke MQTT/intranet dan downlink", "RF ↔ Wi-Fi/IT trust boundary", "Store/queue terbatas; broker reconnect/drop harus terlihat"),
        ("Broker/server", "Message transport, decode, topology, dashboard, dataset", "OT/IT security, persistence, operator access", "Telemetry/alarm/command dapat tertunda atau hilang; recovery harus diuji"),
        ("Pertamina systems", "Network, site power, HSE, asset/permit, operator", "Ownership dan approval", "Tidak ada site acceptance tanpa keputusan tertulis"),
    ],
    [1250, 2400, 3100, 2610],
    font_size=8,
)

add_picture(
    ROOT / "docs" / "manual" / "gateway-intranet-registration-block-diagram.png",
    6.25,
    "Gateway intranet registration reference",
    "Diagram blok referensi yang menunjukkan Gateway Wi-Fi STA, MAC address, broker/server address, dan destination TCP port pada intranet Pertamina.",
    "Gambar 1. Diagram referensi onboarding Gateway. Placeholder MAC/host/port harus diganti; port 1884 pada current source adalah bench, bukan keputusan production.",
)

page_break()

# GLD datasheet
heading("4. Datasheet terkendali: GLD1", 1)
callout("STATUS PRODUK", "Engineering prototype / source-static. Belum ada bukti build current source, measurement hardware, environmental qualification, hazardous-area certificate, atau field acceptance.", PALE_AMBER, AMBER)

heading("4.1 Ringkasan stakeholder", 2)
table(
    ["Atribut", "Nilai yang dapat diterbitkan", "Status/batas"],
    [
        ("Fungsi", "Unit edge multi-gas: membaca sampai 8 kanal MQ melalui front-end analog, menjalankan local processing, alarm/status, dan uplink ke CH.", "CONFIRMED-SOURCE; performa sensing TBD"),
        ("Controller", "ESP32-S3 pada desain/source GLD1.", "Part/source; hardware population TBD"),
        ("Radio", "STAR LoRa 920,0 MHz; BW 125 kHz; SF7; CR 4/5; TX config 17 dBm.", "Current config; bukan RF/regulatory proof"),
        ("Sensor interface", "8 socket MQ; heater setiap socket terhubung langsung ke rail +5 V dan ground.", "Vendor/MPN/lot MQ tidak ditemukan"),
        ("Power", "Input 24 VDC melalui buck ke 5 V dan 3,3 V.", "24 VDC continuous adalah baseline GLD1"),
        ("PSU provisional", "24 VDC/2 A per GLD1, branch-protected.", "PROPOSAL capacity; bukan continuous draw"),
        ("Estimasi board", "Sekitar 0,4–0,8 A pada 24 V setelah margin 25%, tanpa arus alarm/fan eksternal yang belum diketahui.", "DERIVED; wajib ukur"),
        ("Environment", "Outdoor/process-area concept.", "IP/IK, ambient, corrosion, EMC, Ex: TBD/NO-PROOF"),
        ("Mechanical", "Enclosure concept tersedia pada gambar; interface final belum ada.", "Dimensi, weight, material, gland, flamepath, bracket: TBD"),
    ],
    [1900, 4600, 2860],
    font_size=8,
)

heading("4.2 Fakta EasyEDA GLD1", 2)
bullets([
    "Arsip source-GLD_Project.zip SHA-256 A4A6701719C26434998FF3F834F060A5B1FF8916112C7870DD29AEABD1F1E042.",
    "Sheet_1: 928 shapes, 183 komponen; delapan socket MQ U1/U7/U12/U16/U20/U24/U28/U32, heater ke +5 V dan ground.",
    "Sheet_2: 575 shapes, 99 komponen; input/power, ESP32, LoRa, alarm/fan.",
    "Raw PCB: 281 footprint, 940 pad, 217 net, 28 pad tanpa net.",
    "Power utama: CN2/J1 → F1 → Q3 → U36 LMR51450 → +5 V; U43 TPS62162 → VCC/3,3 V.",
    "ULN2003 dan CN1 menyediakan beban fan/alarm. Current source memaksa fan ON terus pada mode 24 V GLD1; arus fan belum tersedia."
])

heading("4.3 Blocker power-tree GLD1", 2)
callout("BLOCKER GLD1-01", "F1 GLD1 diberi part MINISMDC260F/16. Datasheet Littelfuse menilai varian ini 16 V; tidak sesuai untuk dianggap protection input 24 V tanpa penggantian/review. Pilih part berating di atas rail dan transient site, lalu validasi koordinasi fuse/TVS/cable.", PALE_RED, RED)

heading("4.4 Sensor dan batas klaim gas", 2)
p("EasyEDA hanya menyebut MQSocket/header; tidak ada manufacturer, MPN, purchase BOM, lot, calibration certificate, atau foto marking sensor. Nilai Winsen pada laporan ini hanya referensi desain primer untuk sizing—bukan bukti bahwa sensor terpasang adalah Winsen atau varian yang sama.")
table(
    ["Reference candidate", "Heater max menurut vendor", "Batas penggunaan"],
    [
        ("MQ-2", "5 V, ≤950 mW", "Reference only"),
        ("MQ-3B", "5 V, ≤900 mW", "Varian B belum terbukti"),
        ("MQ-4", "5 V, ≤950 mW", "Reference only"),
        ("MQ-5", "5 V, ≤950 mW", "Reference only"),
        ("MQ-6", "5 V, ≤950 mW", "Reference only"),
        ("MQ-7B", "5 V high / 1,5 V low, ≤950 mW", "Varian/siklus heater perlu audit"),
        ("MQ-8", "5 V, ≤900 mW", "Reference only"),
        ("MQ135", "5 V, ≤950 mW", "Reference only"),
    ],
    [2500, 3300, 3560],
    font_size=8,
)
p("MQ bersifat cross-sensitive. Label ‘target gas’ vendor bukan bukti selectivity atau threshold field. Dataset/chamber plan harus mengontrol gas, concentration, distance, humidity, temperature, background, interferents, exposure/recovery, aging, dan reference instrument. MQ-7 memerlukan heater cycle; current power-path yang sekadar mengaktifkan load switch belum membuktikan rail 1,5 V/cycle yang sesuai.")

heading("4.5 Bukti visual GLD1", 2)
add_picture(
    FIELD / "GLD.png",
    2.75,
    "GLD1 enclosure concept",
    "Gambar konsep GLD1 dengan head sensor biru di atas enclosure silinder abu-abu; tidak terlihat Ex marking atau certificate label.",
    "Gambar 2. Konsep bentuk GLD1 dari folder Cilacap. Warna/bentuk tidak membuktikan material, IP/IK, flamepath, gas group, temperature class, EPL, gland, atau sertifikat Ex.",
)

page_break()

# Power decision
heading("5. Keputusan daya GLD1 24 VDC", 1)
heading("5.1 Jawaban praktis", 2)
callout("REKOMENDASI", "Gunakan PSU 24 VDC/2 A per GLD1 sebagai kapasitas nameplate provisional. Konsumsi board diperkirakan sekitar 0,4–0,8 A worst-design setelah margin; 2 A bukan arus kontinu. PSU 1 A hanya dapat dipilih setelah FAT, derating, cable-drop, fan/alarm, dan thermal lulus.", PALE_GREEN, GREEN)

heading("5.2 Rumus dan asumsi", 2)
p("Rumus sizing dasar:")
callout("RUMUS", "I24 = (5 × Iheater,total + Pelektronik,ekuivalen-5V) / (24 × η) + Ibeban-langsung-24V. Laporan memakai allowance elektronik/radio 3 W dan efisiensi buck 0,80–0,90; keduanya asumsi desain, bukan hasil ukur.")
table(
    ["Skenario", "Heater", "Perhitungan input 24 V", "Setelah margin 25%"],
    [
        ("Input user rendah", "8 × 0,1 A @5 V = 4 W", "(4+3)/(24×0,90) = 0,324 A", "0,405 A"),
        ("Input user tinggi", "8 × 0,2 A @5 V = 8 W", "(8+3)/(24×0,80) = 0,573 A", "0,716 A"),
        ("Datasheet steady", "≈7,6 W", "(7,6+3)/(24×0,80) = 0,552 A", "0,69 A"),
        ("Cold design envelope", "≈9 W", "(9+3)/(24×0,80) = 0,625 A", "0,78 A"),
    ],
    [1950, 2500, 2960, 1950],
    font_size=8,
)

p("Delapan heater candidate pada batas 0,9–0,95 W memberi sekitar 7,5–7,6 W steady. Envelope 8–9 W cold-start masih masuk akal dari resistance minimum, tetapi harus ditangkap dengan instrument bandwidth memadai. LMR51450 berating IC 36 V/5 A; rating IC tidak membuktikan PCB, inductor, connector, enclosure, dan thermal sanggup 5 A.")

heading("5.3 Kenapa 2 A tetap rasional", 2)
bullets([
    "Memberi headroom cold heater/inrush dan load-step LoRa/ESP32.",
    "Menanggung derating ambient dan thermal enclosure.",
    "Mengurangi risiko cable/connector voltage drop pada route site yang belum diketahui.",
    "Menyediakan kapasitas untuk fan/alarm GLD1 yang arusnya belum tersedia.",
    "Menjaga operasi di bawah sekitar 80% rating setelah derating sebagai acceptance target provisional.",
    "Memudahkan procurement standardization tanpa menyatakan konsumsi 48 W."
])

heading("5.4 Input daya site yang harus diklarifikasi", 2)
table(
    ["Item", "Kebutuhan keputusan"],
    [
        ("‘110VA → 24VDC’", "Kemungkinan maksudnya 110 VAC, bukan 110 VA. Konfirmasi voltage, frequency, phase, earthing, panel source, available current, breaker, dan UPS/emergency source."),
        ("PSU", "Gunakan industrial isolated PSU dalam lokasi/panel yang disetujui; input/output protection, derating curve, EMC/surge, terminal, certification, dan temperature rating harus masuk BOM."),
        ("Cable", "Butuh length, conductor area, material, ambient, grouping, gland, armor/shield, permissible drop, and short-circuit protection. Hitung Vdrop dua arah dan ukur di terminal J1 saat full load."),
        ("Grounding", "Tentukan PE, functional ground/shield termination, bonding enclosure, surge path, dan separation dengan IS/non-IS circuit bila relevan."),
        ("Battery", "Tidak direkomendasikan untuk GLD1 continuous saat ini karena heater/preheat. Tetap R&D sampai energy profile dan calibration stability terbukti."),
    ],
    [2300, 7060],
    font_size=9,
    first_col_bold=True,
)

heading("5.5 FAT power minimum", 2)
numbers([
    "Catat manufacturer/MPN/lot delapan MQ; ukur resistance heater saat dingin.",
    "Log serempak output PSU, terminal input GLD1, rail +5 V, rail 3,3 V, dan heater tiap socket MQ.",
    "Lakukan ≥20 cold-start setelah OFF ≥30 menit; record 0 s, 1 s, 10 s, 1/5/10/30/60/120 menit.",
    "Stress all-heater, LoRa TX maksimum, RS485, LED, alarm, dan fan bila ada.",
    "Target heater terjauh 4,9–5,1 V termasuk saat TX/alarm; final tolerance mengikuti MPN aktual.",
    "Simulasikan kabel terpanjang/gauge aktual dan closed enclosure pada ambient desain; ukur hotspot.",
    "Pass bila tidak ada current-limit, brownout/reset, rail out-of-limit, dan worst steady ≤80% rating setelah derating; T-class/certificate menetapkan thermal limit final.",
    "Ganti/validasi F1 untuk input 24 V sebelum field acceptance."
])

page_break()

# CH datasheet
heading("6. Datasheet terkendali: CH", 1)
callout("STATUS PRODUK", "CH Ver5 adalah engineering prototype dual-radio dengan charger single-cell dan input SOLAR. Belum ada bukti autonomy, enclosure qualification, RF site test, atau hazardous-area certification.", PALE_AMBER, AMBER)

heading("6.1 Inventaris EasyEDA lengkap", 2)
bullets([
    "CH.zip SHA-256 92B4BB13CE830013413AD3A9D7B4811229A0E3C78ABD5808BA93E7DD18324E6C; berisi satu nested ZIP dualRadioCH_E220Ver5_05852bbac21f4f5f8f88fd60a379ff12.zip.",
    "Nested ZIP memuat 1-Schematic_dualRadioCH_E220Ver5.json, 1-PCB_PCB_dualRadioCH_E220Ver5.json, dan README.txt; EasyEDA editor version 6.5.57.",
    "Schematic satu sheet: 509 raw object dan 71 komponen aktual. PCB: 435 raw object; 71 LIB, 219 TRACK, 141 VIA, 2 COPPERAREA, 2 TEXT.",
    "Outline raw span 83 × 429 EasyEDA units dan tidak boleh dikonversi menjadi mm tanpa Gerber/fabrication drawing. Tidak ada standalone HOLE object pada Ver5. Default DRC raw adalah configuration, bukan bukti DRC pass.",
    "Ver4 memiliki empat mechanical holes; Ver5 memerlukan tray/standoff internal dan positive retention untuk battery/cell yang ditetapkan drawing mekanik."
])

heading("6.2 Controlled datasheet", 2)
table(
    ["Atribut", "Nilai current-source/design", "Status/batas"],
    [
        ("Fungsi", "Menerima GLD1 melalui STAR dan meneruskan ke CH/Gateway melalui MESH.", "CONFIRMED-SOURCE; field capacity TBD"),
        ("MCU", "ESP32-S3-WROOM-1U-N16R8", "EasyEDA component"),
        ("Radio", "2 × E22-900MM22S + 2 SMA feedthrough", "Need dual antenna/cable schedule"),
        ("STAR", "920,0 MHz, BW125, SF7, CR4/5, sync 0x12, TX 17 dBm, preamble 8", "Current config; no site/RF proof"),
        ("MESH", "921,0 MHz, BW125, SF9, CR4/5, sync 0x34, TX 17 dBm, preamble 8", "Current config; no site/RF proof"),
        ("Firmware source", "CH 0.7.3; protocol 0.2.0", "Not proof of flashed version"),
        ("Production hello", "300 s default", "Field-test 30 s + jitter 5 s is temporary; restore before production"),
        ("Power path", "BQ25185 charger/power path; TPS63020 ≈3,28 V; holder 18650; SOLAR input", "Panel/cell/thermal/autonomy TBD"),
        ("Watchdog", "TPL5010; DONE pulse source every 10 s", "Runtime/hardware proof absent"),
        ("Environment/mechanical", "Outdoor/pole concept", "Dimensions, weight, IP/IK, corrosion, wind, Ex, drawing: TBD"),
    ],
    [1800, 4800, 2760],
    font_size=8,
)

heading("6.3 Komponen power utama", 2)
table(
    ["Fungsi", "Component/design", "Interpretasi yang aman"],
    [
        ("Charger/power path", "BQ25185DLHR; R15 18 kΩ; R14 1 kΩ; R16 fixed 10 kΩ", "Single-cell 4,2 V; input limit ≈500 mA; fast charge ≈300 mA; fixed TS means no actual battery temperature sensor"),
        ("3,3 V rail", "TPS63020 + 1,2 µH; feedback 1 MΩ/180 kΩ", "Output sekitar 3,28 V; peak/thermal harus diukur"),
        ("Battery", "BH-18650 holder", "Tidak menetapkan cell maker/model, chemistry, capacity, protection, certification, temperature, cycle life"),
        ("Solar input", "CN3 SOLAR + resettable fuse/diode/TVS/capacitor", "Tidak ada MPPT; input IC 3–18 V tidak berarti sembarang panel aman"),
        ("Load envelope", "ESP32 supply min capability 0,5 A; E22 TX tipikal ≈100 mA/radio", "Verifikasi rail peak sekurangnya 0,7 A sebelum margin bila dual radio aktif; bukan hasil ukur"),
    ],
    [1900, 3100, 4360],
    font_size=8,
)

heading("6.4 Battery dan solar sizing", 2)
p("Tidak ada angka panel watt atau battery Ah yang dapat dibekukan dari ZIP. Gunakan profil daya terukur dan rumus:")
callout("ENERGY MODEL", "Eharian(Wh) = Prata-rata(W) × 24. Battery(Ah) ≥ Eharian × hari-otonomi / (Vnom × usable-DoD × efisiensi). Panel(W) ≥ Eharian / (peak-sun-hours × efisiensi-charge × derating-lapangan).")
bullets([
    "Limit input 500 mA setara maksimum nominal sekitar 2,5 W pada 5 V; battery dapat memasok selisih peak.",
    "Contoh cell 3.000 mAh memerlukan lebih dari 10 jam pada 300 mA sebelum taper dan load CH. Ini contoh hitung, bukan BOM.",
    "Ukur idle, dual RX, masing-masing TX, dual-radio peak, startup/brownout, charge aktif, malam, dan profil 24 jam pada cadence 300 s serta test 30 s+jitter.",
    "Masukkan shading, panel temperature/Voc cold, H₂S/salt/corrosion, rain/dust, cleaning interval, dan target autonomy Pertamina.",
    "Minta IEC 62133-2 dan UN 38.3 untuk sel/pack; keduanya tidak menggantikan Ex certification assembly."
])

heading("6.5 Mounting/antenna implication", 2)
bullets([
    "Dua radio berarti dua jalur 50 Ω yang dilabel STAR 920 dan MESH 921. Satu boss pada gambar konsep belum membuktikan dua feedthrough tersedia.",
    "Panel-surya, CH backplate, dan antenna standoff sebaiknya menjadi modul terpisah agar angle, service, dan cable strain dapat diatur.",
    "Material candidate 316L/passivated dengan galvanic isolation cocok sebagai starting point coastal/chemical, tetapi final bergantung corrosion data site.",
    "Tentukan antenna outdoor, gain/polarization, separation, metal clearance, coax loss/max length, bend radius, UV/chemical jacket, drip loop, strain relief, bonding, dan lightning/surge protection.",
    "Jangan drill/weld aset Pertamina tanpa izin; slotted bracket + band clamp/dual U-bolt dapat menjadi universal concept setelah pole OD/load disurvei."
])

add_picture(
    FIELD / "CH&GW.png",
    2.6,
    "CH and Gateway enclosure concept",
    "Gambar konsep enclosure silinder abu-abu untuk CH dan Gateway dengan satu boss di bagian atas; tidak terlihat Ex marking, gasket, atau dimensional reference.",
    "Gambar 3. Konsep bentuk CH/Gateway. Tidak membuktikan material, ukuran, IP/IK, UV/corrosion, flamepath, jumlah RF feedthrough, atau Ex certificate.",
)

page_break()

# Gateway datasheet
heading("7. Datasheet terkendali: Gateway", 1)
callout("STATUS PRODUK", "Current source membuktikan fungsi bridge radio–MQTT, bukan hardware/network/deployment live. Uplink current source hanya Wi-Fi STA; Ethernet dan WPA-Enterprise/802.1X belum ditemukan.", PALE_AMBER, AMBER)

heading("7.1 Controlled datasheet", 2)
table(
    ["Atribut", "Current source / meeting input", "Status/batas"],
    [
        ("Fungsi", "Menerima MESH dari CH, menerbitkan MQTT/status, menerima downlink, dan mengirim ke radio.", "CONFIRMED-SOURCE"),
        ("Controller", "ESP32-S3; firmware source 0.1.5; protocol 0.2.0.", "Not proof flashed"),
        ("Radio", "SX1262, MESH 921,0 MHz; BW125; SF9; CR4/5; TX config 17 dBm.", "No RF/regulatory field proof"),
        ("IT uplink", "Wi-Fi STA menggunakan WiFi.begin; DHCP adalah inferensi; tidak ditemukan static IP, Ethernet, atau WPA-EAP.", "Current capability"),
        ("Provisioning", "SSID/password, broker host/port/user/password dapat disimpan NVS melalui operator workflow.", "Live security/onboarding not tested"),
        ("MQTT", "WiFiClient plaintext; status period 10 s; RAM queue 8 × 1024 B volatile.", "TLS blocker; queue not durable"),
        ("Status", "MQTT status berisi STA MAC, IP, RSSI, broker, port, queue.", "Pra-onboarding serial status belum menampilkan STA MAC"),
        ("Power/placement", "Meeting note: 5 V indoor/safe area; antenna ke rooftop.", "Current, connector, PSU, dimensions, IP, lightning: TBD"),
        ("Network endpoint", "Default host/SSID masih CHANGE_ME; port 1884 adalah bench.", "Production FQDN/IP/port TBD by IT"),
    ],
    [1800, 4700, 2860],
    font_size=8,
)

heading("7.2 Batas queue dan failure behavior", 2)
p("Queue Gateway hanya 8 item × 1024 byte di RAM; restart/power loss menghilangkan antrean. Ini bukan store-and-forward historian. Acceptance harus menguji disconnect, overflow/drop, reconnect, ordering/replay, alarm handling, dan visibility. Gangguan network tidak boleh menghasilkan ACK alarm palsu; command-result correlation end-to-end masih menjadi residual gap yang harus dibuktikan.")

heading("7.3 Kebutuhan fisik Gateway", 2)
bullets([
    "Lokasi indoor/safe-area tertulis dan accessible; asset label memuat serial serta Wi-Fi STA MAC sebelum onboarding.",
    "PSU 5 V industrial: voltage/current/connector/grounding/UPS/TBD harus diukur dan dibakukan.",
    "Antenna rooftop memerlukan route/coax loss, weatherproof feedthrough, lightning/surge protection, bonding/earthing, structural permission, work-at-height, dan separation dari antenna lain.",
    "Gain/EIRP, EMC, SDPPI/regulatory status, IP/IK, ambient, dimension, weight, dan enclosure certificate belum ditemukan; semua ditandai TBD.",
    "Future Ethernet hanya roadmap bila tidak ada implementation evidence di current source."
])

page_break()

# Server
heading("8. Server, dashboard, database, dan deployment", 1)
heading("8.1 Apa yang repo buktikan", 2)
bullets([
    "Ada Node-RED bridge untuk MQTT input/output, decode, topology/status, command, dan dashboard-related flows.",
    "Ada dataset recorder dengan MySQL + CSV/idempotency. MySQL pada flow ini adalah recorder dataset, bukan bukti historian telemetry production sudah dirancang.",
    "Secret seharusnya berasal dari runtime environment/credentials; generator non-loopback menuntut TLS dan credential.",
    "Flow snapshot utama saat audit berisi 56 node dan tidak memiliki ingest/decode gld/gateway/status yang sudah ada pada generator terkini. Read-only check melaporkan drift:true. Dataset-flow check melaporkan Drift:false.",
    "Tidak ditemukan manifest production authoritative untuk OS, Node.js, Node-RED/plugin version, broker package/config/service, package.json/container/service, reverse proxy, backup/restore, retention, HA, RPO/RTO, monitoring, atau dashboard HTTPS/RBAC."
])

heading("8.2 Blocker server flow drift", 2)
callout("SERVER-01", "Rekonsiliasi generator dan pertamina-gld-server.flow.json sebelum packaging/deployment. Setelah source snapshot konsisten, verifikasi revision/hash yang benar-benar live; audit sesi ini tidak mengakses deployment live.", PALE_RED, RED)

heading("8.3 Baseline pilot sementara", 2)
table(
    ["Item", "Minimum planning", "Recommended pilot planning", "Finalization input"],
    [
        ("CPU", "2 vCPU", "4 vCPU", "Node count, decode/flow load, concurrent users, alarm burst"),
        ("RAM", "4 GB", "8 GB", "Node-RED heap, broker, DB cache, OS/monitoring"),
        ("Storage", "50 GB SSD", "100 GB SSD", "Retention, DB/index growth, logs, backup, evidence images"),
        ("Network", "1 GbE/VLAN", "1 GbE + managed firewall path", "IT architecture, HA, backup network"),
        ("OS", "Server/LTS approved IT", "Windows Server 2022 or Linux LTS after qualification", "Current scripts are Windows-oriented; standard site platform wins"),
        ("Availability", "Single pilot", "Service auto-start + health + backup", "RPO/RTO, HA policy, maintenance window"),
    ],
    [1250, 1800, 2700, 3610],
    font_size=8,
)
p("Semua angka di atas adalah PROPOSAL/baseline planning, bukan minimum vendor yang telah diuji. Windows Server 2022 cenderung menurunkan migration risk untuk PowerShell/current path assumptions; Linux LTS tetap mungkin setelah scripts, service, path, permissions, plugin, dan backup diuji.")

heading("8.4 Storage planning bound", 2)
p("Bila digunakan batas konservatif 1 KB per GLD1 setiap 10 s untuk 3 GLD1, maka 3 × 8.640 = 25.920 record/hari, sekitar 25,9 MB/hari raw, atau 9,46 GB/tahun raw. Angka ini belum memasukkan protocol overhead, DB/index, logs, replicas, backup, alarm burst, atau dashboard cache—dan juga mungkin jauh di atas payload aktual. Final sizing memakai capture lapangan serta retention policy.")

heading("8.5 Production deployment package yang belum ada", 2)
bullets([
    "Pinned OS/runtime/plugin/broker versions dan checksum package.",
    "Service account, filesystem paths, least privilege, secret storage, certificate store, time synchronization.",
    "Broker auth/TLS/ACL, connection/session policy, retained messages, last-will, limit, logging, and rotation.",
    "Reverse proxy HTTPS, certificate lifecycle, RBAC/session timeout, admin network restriction.",
    "Telemetry historian/retention design terpisah dari dataset recorder; schema/version/migration.",
    "Backup/restore untuk flows, credentials, broker config, replay-state JSON, MySQL dataset, certificates, dan audit logs.",
    "Monitoring/alerting untuk process, disk, DB, broker connections, queue/drop, stale node, certificate expiry, backup failure.",
    "RPO/RTO, HA/DR, patching, vulnerability process, deployment/change/rollback runbook."
])

page_break()

# Network and security
heading("9. Desain intranet, onboarding, dan keamanan", 1)
heading("9.1 Prinsip jaringan yang benar", 2)
bullets([
    "Gateway dan server tidak harus satu SSID atau subnet. Mereka harus memiliki routed reachability yang disetujui dari VLAN Wi-Fi Gateway ke endpoint broker.",
    "Gateway memulai koneksi outbound; server tidak perlu membuka koneksi baru langsung ke Gateway.",
    "Jarak fisik sekitar 2 km tidak menentukan MQTT bila intranet menyediakan routing, latency, loss, firewall, dan DNS yang memadai. Jarak RF CH–GW tetap harus diuji terpisah.",
    "Broker sebaiknya mempunyai FQDN dengan IP statis/reservasi dan lifecycle change yang terkendali. Gateway dapat DHCP jika reservation/onboarding policy mengizinkan.",
    "TCP reachability hanya precheck; MQTT success memerlukan broker evidence/CONNACK sukses, active subscription, publish/receive, auth/ACL, reconnect, dan log."
])

heading("9.2 Data wajib untuk tiket IT", 2)
table(
    ["Data", "Isi sekarang", "Yang harus diisi sebelum submit"],
    [
        ("Wi-Fi STA MAC", "Current source dapat publish MAC setelah online", "Actual MAC per unit pada label/commissioning output sebelum onboarding"),
        ("Broker/server endpoint", "CHANGE_ME_MQTT_HOST", "Production FQDN dan IP/reservation owner"),
        ("Destination port", "1884 bench", "Final TCP port; TLS policy/certificate; protocol MQTT"),
        ("Source network", "Wi-Fi STA/DHCP inference", "SSID, VLAN/subnet, DHCP/NAC/asset policy"),
        ("DNS/NTP", "TBD", "Resolver/port and time source required for TLS/log correlation"),
        ("Identity", "Username/password provisionable", "Per-device identity; rotation/revocation; topic ACL"),
        ("Owner", "TBD", "LGU device PIC, Pertamina IT PIC, operations/HSE PIC, location/serial"),
    ],
    [2100, 3000, 4260],
    font_size=8,
)

heading("9.3 Firewall/port matrix provisional", 2)
table(
    ["Source", "Destination", "Port", "Purpose", "Policy"],
    [
        ("Gateway VLAN", "Broker FQDN/IP", "TBD TCP (1884 bench)", "MQTT uplink/downlink", "Allow only registered Gateway; TLS/auth/ACL decision"),
        ("Gateway VLAN", "DNS", "53 UDP/TCP if needed", "Resolve broker", "Site resolver only"),
        ("Gateway VLAN", "NTP", "123 UDP if needed", "Time/TLS/logs", "Site NTP only"),
        ("Operator/admin", "Reverse proxy/dashboard", "443 TCP proposal", "Dashboard", "HTTPS + RBAC; not from Gateway VLAN"),
        ("Admin subnet", "Node-RED admin", "1880 TCP if retained", "Restricted administration", "Do not expose broadly; reverse proxy/VPN"),
        ("Application server", "MySQL dataset", "3306 TCP if separate", "Dataset recording", "Local/internal only; never from Gateway VLAN"),
        ("Any", "Internet", "None by default", "No default requirement", "Default-deny; exception via change approval"),
    ],
    [1500, 1900, 1350, 2200, 2410],
    font_size=8,
)

heading("9.4 TLS gap dan pilihan keputusan", 2)
callout("SECURITY-01", "Gateway current source memakai WiFiClient plaintext, sedangkan generator server menolak broker non-loopback tanpa TLS dan credential. Karena itu claim ‘TLS end-to-end ready’ tidak benar untuk current source.", PALE_RED, RED)
table(
    ["Opsi", "Kelebihan", "Konsekuensi", "Rekomendasi"],
    [
        ("TLS end-to-end di Gateway", "Auth/integrity/privacy on route; sejalan dengan server policy", "Butuh CA trust, time sync, credential/cert per device, memory/performance, rotation/revocation test", "Utama"),
        ("Exception OT terisolasi", "Dapat mempercepat pilot", "Plaintext tetap ada; wajib VLAN/NAC/firewall/ACL/logging dan approval/risk acceptance", "Sementara, formal only"),
        ("Local bridge plaintext→TLS", "Membatasi plaintext segment", "Menambah asset, power, patching, monitoring, HA, dan another trust boundary", "Transitional fallback"),
    ],
    [2150, 2500, 3210, 1500],
    font_size=8,
)
p("MAC allowlist/NAC membantu asset control, tetapi MAC dapat disalin dan tidak menggantikan authentication/authorization. Jika Pertamina mewajibkan WPA2-Enterprise/802.1X, current Gateway source belum membuktikan dukungan; perlu firmware change atau SSID/bridge exception yang disetujui.")

heading("9.5 Security acceptance", 2)
numbers([
    "Approve architecture, asset inventory, MAC/FQDN/port/VLAN/DHCP/NAC/firewall/DNS/NTP/owner.",
    "Test valid and invalid identity, unauthorized topic/command, certificate trust/expiry/revocation path, and audit log.",
    "Verify dashboard via HTTPS/RBAC/session control; restrict Node-RED admin and DB.",
    "Test Gateway/broker/Node-RED restart, queue/drop/reconnect, replay state, alarm burst, and no false ACK.",
    "Perform 24–72 hour soak with evidence: timestamp, revision/hash, broker log, screenshot, negative tests, PIC, sign-off.",
    "Use NIST OT security principles for segmentation, least privilege, logging, recovery, and change control; apply MQTT authentication/authorization/secure communication guidance."
])

page_break()

# Mounting
heading("10. Mounting, antena, kabel, grounding, dan site interface", 1)
heading("10.1 Prinsip universal bracket", 2)
callout("PROPOSAL MEKANIK", "Gunakan modular backplate terpisah untuk GLD1/CH, interchangeable clamp untuk round pole/handrail, wall/base plate untuk structure-approved anchor, serta bracket terpisah untuk solar panel dan antenna. Jangan freeze ukuran sebelum survey dimensional dan load approval.")
bullets([
    "Candidate material: SS316L/passivated untuk coastal/chemical; gunakan galvanic isolation terhadap carbon steel/aluminium dan validasi corrosion/H₂S/salt-mist site.",
    "Slotted bracket + dual U-bolt/band clamp memberi adjustment untuk beberapa OD; wall/building memakai base plate/anchor yang disetujui structural owner.",
    "M8/M10/M12 dan spanner 13/17/19 mm hanya candidate standardization. Final bolt class, nut, washer, captive/locking method, torque, anti-seize, tool restriction, dan inspection interval harus disepakati.",
    "Tidak ada drilling, welding, hot-work, atau piggyback pada existing detector/support tanpa permit dan asset owner approval.",
    "Mechanical drawing harus memuat mass/CG, wind/seismic load, vibration, drain/orientation, service envelope, removal path, cable bend radius, gland location, ingress/drip loop, earthing point, and tag plate.",
    "GLD1 placement dekat existing sensor/power dapat mengurangi cable, tetapi tidak boleh mengganggu detector, calibration access, airflow, heat, or maintenance."
])
p("Rekaman meeting menyebut extension transmitter sekitar 1,5 m dan pole CH/repeater sekitar 2–2,5 m sekaligus support solar, serta tiga titik solar. Angka ini hanya proposal meeting; jangan dimasukkan ke fabrication drawing sampai actual pole, load, wind, access, antenna clearance, and hazardous boundary disurvei dan disetujui.")

heading("10.2 Tidak ada tinggi/distance universal", 2)
p("Ketinggian GLD1 tidak boleh ditentukan hanya dari nama gas. Densitas/buoyancy dipengaruhi temperature, pressure, release momentum, ventilation, obstruction, weather, and process geometry. Pertamina/HSE harus menyediakan hazard study/area drawing dan target release scenario; vendor kemudian memvalidasi coverage melalui dispersion/risk basis dan gas-release/chamber/site test. Catatan ‘kebocoran hampir harian’ bila benar harus diverifikasi dari incident/operations data—bukan alasan menempatkan prototype uncertified lebih dekat ke leak source.")

heading("10.3 Data survey mekanik wajib", 2)
table(
    ["Domain", "Data minimum"],
    [
        ("Pole/handrail", "OD range, material, wall thickness, height, existing load, allowable load, vibration, corrosion, clearance"),
        ("Structure/wall", "Anchor policy, concrete/steel condition, prohibited zones, drilling/welding/hot-work rules, structural sign-off"),
        ("Environment", "Ambient min/max, solar load, rain, flood, dust, salt, H₂S/chemical, wind/seismic, lightning class"),
        ("Access", "Work-at-height, ladder/scaffold/lifting, maintenance route, emergency egress, tool restriction, permit/JSA"),
        ("Cable", "Source panel, route/length, tray/conduit, cable type/area, gland/armor/shield, separation, fire rating, tagging"),
        ("Antenna", "Height, LOS/obstruction, separation, metal clearance, coax length/loss, rooftop access, grounding, lightning zone"),
        ("Existing detector", "Owner, no-interference distance, calibration envelope, mounting prohibition, power sharing approval"),
    ],
    [2100, 7260],
    font_size=9,
)

heading("10.4 Visual evidence dan lokasi konsep", 2)
add_picture(
    FIELD / "WhatsApp Image 2026-08-18 at 16.01.54.jpeg",
    6.3,
    "RU IV plot map field note",
    "Foto plot map RU IV dengan anotasi kandidat posisi GLD1 A1 A2 A3, CH A4 A5, Gateway indoor dan antenna rooftop, serta staged gas detection plan.",
    "Gambar 4. Plot map historical survey 18 Desember 2025/catatan lapangan. Ini input phase discussion, bukan layout aktif, approved plot plan, area-classification drawing, cable route, RF survey, atau installation permit.",
)

page_break()

# Hazardous area
heading("11. Hazardous area, ATEX/IECEx, dan batas trial", 1)
callout("SAFETY HOLD", "Tidak ditemukan certificate/marking ATEX atau IECEx untuk assembly GLD1 maupun CH. Hingga konfigurasi final lengkap disertifikasi dan diterima Pertamina, status yang benar adalah: ‘Engineering prototype — not certified for installation in a hazardous area.’", PALE_RED, RED)

heading("11.1 Apa arti ATEX/IECEx dalam konteks ini", 2)
bullets([
    "ATEX Directive 2014/34/EU mengatur equipment/protective systems/components untuk potentially explosive atmospheres di pasar EU; istilah ATEX bukan sinonim umum untuk ‘explosion proof’ dan belum otomatis menjadi satu-satunya acceptance basis Indonesia/site Pertamina.",
    "IECEx Certified Equipment Scheme memakai independent testing dan quality assessment. Ex component certificate masih memerlukan additional certification ketika diintegrasikan ke equipment/assembly.",
    "Area classification gas mengacu pada IEC 60079-10-1; selection/installation/documentation/initial inspection pada IEC 60079-14. Protection concept dapat melibatkan flameproof ‘d’, intrinsic safety ‘i’, atau lain-lain—harus dipilih oleh competent Ex design/certification route.",
    "SNI IEC 60079 editions dan ketentuan Migas/Pertamina yang diterima harus dikonfirmasi. Standard edition, certificate scheme, local approval, dan site procedure tidak boleh diasumsikan.",
    "IP/IK, material metal, physical partition, atau enclosure yang tampak industrial tidak sama dengan Ex certification."
])

heading("11.2 Scope yang harus ikut sertifikat/assessment", 2)
table(
    ["Subassembly", "Pertanyaan sertifikasi"],
    [
        ("GLD1 sensor head", "Gas path, sensing element, hot surface, flame arrestor/flamepath, material, calibration access, contamination"),
        ("GLD1 electronics/power", "Protection concept, component ratings, spacing, thermal/T-class, input protection, fault analysis"),
        ("CH electronics", "Dual radio, charger, battery fault/thermal, enclosure, antenna feedthrough"),
        ("Battery/solar", "Cell/pack, protection, cable, panel, charger, connector, temperature, fault energy; whether inside certificate scope"),
        ("Enclosure/gland/seal", "Exact manufacturer/model, certificate schedule, special conditions, torque, cable range, blanking, ingress, flamepath"),
        ("Antenna/surge/earth", "RF output, antenna/cable/gland/surge/barrier configuration, bonding and installation conditions"),
        ("Final assembly", "Marking, gas group, T-class, EPL/category, ambient, certificate number, QAR/QAN, drawings/BOM revision"),
        ("Installation", "Area drawing, equipment selection, cable/earthing, initial inspection by competent person, dossier and maintenance"),
    ],
    [2300, 7060],
    font_size=9,
)

heading("11.3 Pemisahan sensor dan power", 2)
p("Pemisahan electronics/power dari sensor dapat menjadi design direction, tetapi partisi biasa tidak cukup. Separation, seal/gland, flamepath, barrier, creepage/clearance, fault energy, thermal path, and every external connection harus berada dalam protection concept dan certificate/drawing scope. Memindahkan trial dari LPG tank ke SRU juga tidak otomatis aman; SRU/H₂S tetap memerlukan area-classification drawing dan keputusan tertulis.")

heading("11.4 Hold points", 2)
numbers([
    "Pertamina menerbitkan/menyetujui area-classification drawing untuk setiap titik: Zone/non-hazardous, gas group, T-class, EPL/category, gas, ambient, 3D boundary.",
    "LGU/vendor memilih protection concept dan certification basis/edition yang diterima Pertamina/Migas/local authority.",
    "Complete manufacturer BOM, controlled drawings, enclosure/gland/antenna/battery/solar schedule, risk/fault/thermal analysis selesai.",
    "ExCB/laboratory application, quotation, test plan, schedule, certificate scope, QAR/QAN, and special conditions tersedia. Klaim ‘3 minggu’ baru boleh digunakan setelah bukti ini.",
    "Certificate authenticity/status diverifikasi pada controlled database; final marking cocok dengan as-built serial/revision.",
    "Installation dossier, competency, permit/JSA, cable/earthing, inspection record, and maintenance plan disetujui.",
    "Sebelum semua hold point classified ditutup, trial hanya pada lokasi non-hazardous yang ditetapkan Pertamina secara tertulis dan tetap mengikuti permit serta safety control."
])

page_break()

# Site phases and schedule
heading("12. Lokasi kandidat, fase, jadwal, dan pembagian kerja", 1)
heading("12.1 Status lokasi", 2)
table(
    ["Lokasi/fase dari catatan", "Gas/quantity", "Status keputusan"],
    [
        ("OM47 tahap 1", "LPG, 3 sensor", "Printed plot input; handwritten A1/A2/A3 GLD1 dan A4/A5 CH; belum freeze"),
        ("OM47 tahap 2", "LPG, 3 sensor", "Roadmap input; hazardous/site approval TBD"),
        ("KPC tahap 3", "H₂/butane/ethane, 11", "Roadmap input; scope/contract TBD"),
        ("FOC II tahap 4", "LPG/butane/ethane/propane, 6", "Roadmap input; scope/contract TBD"),
        ("SRU/PSU candidate", "H₂S/NH₃ mentioned", "Relocation/trial discussion; not automatically safe; classification TBD"),
        ("Sampit candidate", "LPG/methane/H₂S mentioned", "Meeting note only; site/scope/date TBD"),
    ],
    [2650, 2500, 4210],
    font_size=9,
)

heading("12.2 Rencana kerja berbasis tanggal", 2)
table(
    ["Tanggal target", "Deliverable", "Owner utama", "Exit criteria"],
    [
        ("20 Aug 2026", "Internal decision pack", "LGU", "Enam keputusan disetujui; P0/P1 owner ditetapkan"),
        ("21–25 Aug", "RFI terkonsolidasi", "LGU → Pertamina/vendor", "Network/HSE/power/mechanical/server/contract questions terkirim"),
        ("≤28 Aug", "Datasheet Rev A + measurement plan", "GLD1/CH/GW/server leads", "Controlled fields, TBD register, FAT instruments/samples"),
        ("1–11 Sep", "Pre-survey pack", "LGU + Pertamina", "Point list, drawing request, permit/JSA, survey template, IT ticket draft"),
        ("Next site visit", "Dimensional/HSE/power/RF/network survey", "Joint team", "Signed measurements, photos, routes, classification, RF/network evidence"),
        ("≤18 Sep", "Freeze phase-1 interface", "LGU + Pertamina", "Location, pole, cable, power, antenna, network, owner approved"),
        ("≤25 Sep", "Support/cable/pole submission", "LGU mechanical/electrical", "Drawing/BoQ before end-Sep target; assumptions closed"),
        ("Aug–Dec 2026", "Certification application/workstream", "LGU/vendor/ExCB", "ExCB/application/test schedule evidenced; release date remains TBD; end-2026 is aspiration only"),
        ("Oct 2026 onward", "Design closure + FAT", "LGU/vendor", "P0 closed; controlled BOM/drawing; FAT milestones baselined"),
        ("Before Sep 2027", "SAT readiness for 3 GLD1 + CH", "Joint program", "Contract/scope, certificate, network, installation, training, spares, FAT pass"),
        ("Sep 2027 target", "Phase-1 field acceptance", "Pertamina signatory", "SAT evidence + punch-list closure; target must be confirmed"),
    ],
    [1500, 3300, 1750, 2810],
    font_size=8,
)

heading("12.3 Pembagian tanggung jawab", 2)
table(
    ["Deliverable", "LGU", "Pertamina RU IV / IT / HSE", "Vendor/ExCB"],
    [
        ("Controlled device datasheets", "R/A: content, BOM, revision, evidence status", "C: site requirement/approval", "C: vendor/certificate data"),
        ("Area classification & acceptance basis", "C: design input", "R/A: site drawing, gas/T-class/EPL, procedure", "C: certification route"),
        ("GLD1 power design/FAT", "R/A: protection, sizing, fan/alarm, test", "C: supply/cable/permit", "C: PSU/part data"),
        ("CH energy/mechanical", "R/A: panel/cell/profile/bracket", "C: autonomy, pole/load, environment", "C: battery/solar/enclosure certificates"),
        ("Gateway/network", "R: MAC/device/config/test support", "A: VLAN/NAC/firewall/FQDN/port/TLS policy", "C: antenna/surge"),
        ("Server deployment", "R: package/flow/test/runbook", "A/C: platform/security/operations", "C: platform support"),
        ("Certification", "R: controlled design dossier and samples", "A/C: accepted scheme/site interface", "R/A: assessment/test/certificate in own scope"),
        ("Installation/SAT", "R: commissioning/test evidence", "A: permit, execution owner, sign-off", "C: special conditions/inspection support"),
    ],
    [2100, 2700, 3060, 1500],
    font_size=8,
)
p("R = Responsible, A = Accountable, C = Consulted. Nama orang, kontrak boundary, procurement owner, and signatory masih TBD dan harus menggantikan role generik sebelum distribusi eksternal.", style="Small Note")

page_break()

# RFI
heading("13. RFI/pernyataan yang harus ditutup", 1)
heading("13.1 Kepada Pertamina HSE/site", 2)
numbers([
    "Berikan area-classification drawing untuk A1–A5, OM47, SRU/PSU, dan candidate lain: Zone/non-hazardous, gas/dust group, temperature class, EPL/category, actual gas, ambient, dan 3D boundary.",
    "Tentukan scheme/standard edition yang diterima: IECEx/ATEX/SNI IEC 60079/Migas/internal; minimum IP/IK, corrosion/H₂S/salt-mist, EMC, and signatory.",
    "Apakah battery, solar panel, gland, antenna, surge protector, cable, and complete assembly wajib dalam certificate scope?",
    "Konfirmasi incident/leak basis, target detection scenario, existing detector interface, release/dispersion information, and permitted non-hazardous trial location.",
    "Tetapkan permit/JSA, work-at-height, hot-work/drilling/welding restriction, gas testing, competent Ex personnel, initial inspection, dan maintenance dossier.",
    "Konfirmasi window survey beberapa hari dan alokasi sekitar satu hari untuk induction/permit/JSA sebelum akses SRU: personel pendamping, contractor, equipment entry, dan pembagian agenda dimensional/RF/power/network."
])

heading("13.2 Kepada Pertamina electrical/mechanical", 2)
numbers([
    "Konfirmasi apakah source ‘110VA’ sebenarnya 110 VAC; berikan voltage/frequency/phase/earthing, panel, breaker, spare capacity, UPS/emergency status, dan cable route.",
    "Berikan pole/handrail/building OD/material/thickness/height/allowable load/wind/seismic/vibration/corrosion serta anchor policy.",
    "Setujui allowable bolt/nut/washer/material/torque/anti-loosening/tool size dan inspection interval.",
    "Berikan cable standard, conductor/armor/shield/gland/termination/tagging/fire rating/segregation, allowable drop, and grounding/bonding/lightning rules.",
    "Konfirmasi antenna rooftop access, height/clearance, lightning zone, coax route/max length, grounding, work-at-height, and structural approval.",
    "Tetapkan target autonomy CH, availability, maintenance access, panel cleaning interval, shading/solar data, dan allowed downtime."
])

heading("13.3 Kepada Pertamina IT", 2)
numbers([
    "SSID/band dan authentication: PSK atau 802.1X; captive portal dilarang/ditangani bagaimana?",
    "VLAN/subnet, DHCP reservation/static policy, gateway, DNS, NTP, NAC/MAC allowlist, asset registration, and location/owner fields.",
    "Broker FQDN/IP, final TCP port, TLS/mTLS/VPN policy, CA/certificate issuance/rotation/revocation, per-device identity, and topic ACL.",
    "Firewall matrix, internet default-deny, logging/SIEM/retention, connection monitoring, incident response, and change window.",
    "Approved server platform: OS/version, VM/physical, CPU/RAM/storage, network zone, backup, monitoring, patching, RPO/RTO/HA.",
    "Dashboard HTTPS/RBAC/SSO/session policy; admin subnet; Node-RED and DB exposure restrictions.",
    "Acceptance evidence format: CONNACK/broker logs, negative tests, packet capture authorization, screenshot, timestamp/NTP, PIC, and sign-off."
])

heading("13.4 Kepada vendor/desain", 2)
numbers([
    "Complete manufacturer BOM untuk GLD1/CH/GW: exact part/model/lot termasuk MQ, fan/alarm, PSU, battery, panel, enclosure, glands, antenna/coax/surge/bracket.",
    "Gerber, fabrication/assembly drawing dalam mm, hole table, STEP/3D, mass/CG, enclosure drawing, gasket/gland/torque/cable schedule.",
    "GLD1: selesaikan F1 untuk input 24 V, MQ-7 heater cycle, input transient/TVS/thermal, fan/alarm current, dan identitas sensor per socket.",
    "CH: exact cell/pack and panel data; compatibility BQ25185; 24 h current profile; autonomy/recharge/brownout/dual-radio/thermal test; IEC 62133-2/UN 38.3 evidence.",
    "RF: antenna/coax/gland/surge datasheet, gain/EIRP calculation, dual-radio isolation, link budget, and site acceptance criteria RSSI/SNR/packet-loss.",
    "If Ex is claimed: CoC/DoC/QAR/QAN, certificate schedule, special conditions, marking, ambient, gas group, T-class, EPL, drawing/BOM revision, and online verification record."
])

heading("13.5 Scope/contract questions", 2)
numbers([
    "Apakah phase 1 benar 3 GLD1 + 1 CH, dan apakah Gateway/server termasuk supply/installation/operation scope?",
    "Apakah additional contract diperlukan untuk site survey, pole/cable/support, certification, server/network, FAT/SAT, training, spares, and maintenance?",
    "Siapa accountable owner dan signatory untuk HSE, IT, electrical, mechanical, operations, procurement, cybersecurity, certification, FAT, and SAT?",
    "Konfirmasi target end-September 2026 dan September 2027 beserta dependency, review duration, approval lead time, and change control.",
    "Tentukan definition of done, warranty/support, spare strategy, calibration interval, data ownership/retention, and post-SAT operational handover.",
    "Kapan memo/rekomendasi Pertamina yang menjadi dasar formal scope diterbitkan, dan role mana yang berwenang menyetujui lokasi sebelum contractor serta IT dilibatkan?"
])

page_break()

# Acceptance
heading("14. FAT, SAT, dan bukti penerimaan", 1)
heading("14.1 FAT matrix", 2)
table(
    ["Domain", "Test minimum", "Pass evidence"],
    [
        ("GLD1 power", "20 cold-start; 120 min profile; all heaters; LoRa/RS485/LED/alarm/fan; cable drop; +5 V/3,3 V; closed-enclosure thermal", "Waveform/log calibrated instrument; rail within MPN limit; no reset/current-limit; ≥20% headroom; blocker closed"),
        ("GLD1 sensing", "MPN/lot; heater resistance/cycle; channel map; zero/span; target/interferent; humidity/temp; exposure/recovery; reference instrument", "Calibration/dataset report; confusion/false alarm/miss criteria; traceability"),
        ("CH energy", "24 h production/test cadence; idle/RX/TX/dual peak; charge/day/night; autonomy; recharge; brownout; temp", "Energy budget; measured profile; cell/panel BOM; autonomy/recovery pass"),
        ("Mechanical", "Fit/retention; fastener torque; vibration; pull/load; ingress; corrosion basis; cable/gland; 2 RF feedthrough", "Drawing/revision, photos, torque sheet, inspection/test report"),
        ("RF", "Bench protocol; range topology; interference; packet/sequence/retry; dual-radio isolation", "RSSI/SNR/PER/latency capture; no unaccounted loss"),
        ("GW/server", "Provisioning, auth/TLS/ACL, normal/alarm/downlink, disconnect/restart/replay, 24–72 h soak, dashboard/security/backup-restore", "Broker/server logs, flow hash, screenshot, negative test, restore proof"),
        ("Hazardous", "Certification test/assessment per agreed scheme", "Valid certificate/QAR/marking/drawing/BOM special conditions; not self-declared"),
    ],
    [1450, 4950, 2960],
    font_size=8,
)

heading("14.2 SAT matrix", 2)
table(
    ["Stage", "Checks", "Evidence/sign-off"],
    [
        ("Pre-install", "Approved drawing/BOM/certificate; serial/marking; permits/JSA; classification; cable/pole/power/network readiness", "Release note and hold-point signatures"),
        ("Mechanical/electrical", "Mount/orientation/torque/load/access; gland/drip/ground/bond/surge; 24 V polarity/no-load/full-load/drop; CH panel/battery", "As-built photos/drawing, torque/power sheet"),
        ("RF/network", "Actual height/location; RSSI/SNR/PER; Gateway IP/DNS/NTP; TCP precheck; MQTT CONNACK/subscriptions/auth/ACL", "RF capture, broker log, IT sign-off"),
        ("End-to-end", "GLD1→CH→GW→broker→Node-RED→dashboard; alarm; downlink/result; timestamp/sequence; disconnect/recovery", "Correlated event pack; no false ACK"),
        ("Safety", "Certificate/config match; special conditions; initial Ex inspection; surface temperature/ground; no interference existing detector", "Competent inspector record/HSE sign-off"),
        ("Soak/handover", "24–72 h; stale/drop/reconnect; backup/restore; operator training; spares/calibration/maintenance", "Acceptance report, punch list, owners/dates"),
    ],
    [1600, 5000, 2760],
    font_size=8,
)

heading("14.3 Evidence package per test", 2)
bullets([
    "Test ID, requirement/source, device serial, HW/FW/flow/config revision dan hash.",
    "Date/time/timezone, synchronized clock source, location/zone, environmental condition.",
    "Instrument maker/model/serial/calibration status; setup drawing and cable/antenna/power details.",
    "Raw log/capture, calculation, screenshot/photo, expected/actual, pass/fail, anomaly, retest reference.",
    "PIC executor/reviewer, Pertamina witness/signatory, open punch item, owner, due date.",
    "Explicit separation: static-source evidence, build evidence, hardware measurement, network/MQTT evidence, deployment revision, and field acceptance."
])

heading("14.4 Definition of readiness", 2)
table(
    ["Gate", "Ready bila"],
    [
        ("Ready for internal design review", "This report approved; owners/dates assigned; source/BOM/drawings controlled; P0 plan funded"),
        ("Ready for non-hazardous pilot", "Pertamina written non-hazardous point; permits; FAT power/RF/network/server pass; mechanical/electrical survey closed"),
        ("Ready for classified-area installation", "Area classification + accepted complete-assembly certification + installation design + competent initial inspection + HSE/IT/site approvals"),
        ("Ready for production handover", "SAT/soak/backup/restore/security/training/spares/calibration/maintenance/ownership pass; punch list controlled"),
    ],
    [2800, 6560],
    font_size=9,
)

page_break()

# Sources
heading("15. Daftar sumber dan register bukti", 1)
heading("15.1 Sumber lokal primer", 2)
table(
    ["ID", "Sumber", "Penggunaan dan batas"],
    [
        ("L-01", r"C:\Users\MSI\Downloads\Summary recorder 081426 .pdf", "Notulen/meeting input; bukan sertifikat atau acceptance"),
        ("L-02", r"C:\Users\MSI\Downloads\20260814_092122.m4a", "Audio meeting 02:15:16,693; offline ASR + cross-check; ucapan adalah meeting evidence, bukan instruksi atau certificate"),
        ("L-03", r"C:\Users\MSI\Downloads\cilacap\ (7 gambar unik)", "Konsep enclosure, plot map, catatan kebutuhan; bukan dimensional/Ex/RF proof"),
        ("L-04", r"C:\Users\MSI\Downloads\CH.zip", "Nested EasyEDA schematic/PCB inventory; static design only"),
        ("L-05", r"docs\wiring\gld-project-ver2-2026-07-01\source-GLD_Project.zip", "GLD1 full schematic/PCB static"),
        ("L-06", r"firmware\config\*.h; firmware\gld/ch/gateway\src; platformio.ini", "Current source configuration/behavior; not build/flash/runtime proof"),
        ("L-07", r"server\nodered\ generator/snapshot/functions/README", "Server source and read-only drift result; not live deployment proof"),
        ("L-08", r"docs\manual and output\pdf old architecture/task documents", "Cross-check only; new report restates all needed content"),
    ],
    [850, 4300, 4210],
    font_size=8,
)

heading("15.2 Current-source evidence pointers", 2)
table(
    ["Area", "Path/line pointer"],
    [
        ("GLD1 power/profile", r"firmware\gld\src\GldUnifiedMain.cpp:5428; source-GLD_Project.zip full EasyEDA schematic/PCB"),
        ("CH radio/power/runtime", r"LoraStarConfig.h:13–49; LoraMeshConfig.h:13–48; ChConfig.h:35–40,77–92; ChStarMeshRuntimeMain.cpp:360–383,1054–1065"),
        ("CH active profile", r"firmware\platformio.ini:287–309,332–337; ChBoardPinsCh3.h:26–31"),
        ("GW Wi-Fi/MQTT", r"GatewayMqttMeshMain.cpp:238–241,636–655,797–870,1589–1686,1759–1763"),
        ("GW config", r"ServerConfig.h:36–64; GwConfig.h:19–35; FirmwareVersion.h:5–14; platformio.ini:363–380"),
        ("Server security/drift", r"apply-pertamina-gld-flow.js:105–116,2212–2224,2272–2288; pertamina-gld-decode.js:1023–1130"),
        ("Server capability", r"server\nodered\README.md:1–98; apply-pertamina-gld-dataset-flow.ps1:1–16,32–53"),
    ],
    [2300, 7060],
    font_size=8,
)

heading("15.3 Sumber primer vendor/standard", 2)
official_sources = [
    ("R-01", "ATEX Directive 2014/34/EU", "https://eur-lex.europa.eu/eli/dir/2014/34"),
    ("R-02", "Workplace explosive atmospheres Directive 1999/92/EC", "https://eur-lex.europa.eu/eli/dir/1999/92/oj"),
    ("R-03", "IECEx Certified Equipment Scheme overview", "https://www.iecex.com/certified-equipment-scheme/overview/"),
    ("R-04", "IECEx FAQ / certificate verification guidance", "https://www.iecex.com/resources-and-news/frequently-asked-questions/"),
    ("R-05", "IEC 60079-10-1:2020 area classification", "https://webstore.iec.ch/en/publication/63327"),
    ("R-06", "IEC 60079-14:2024 design, selection, installation, initial inspection", "https://webstore.iec.ch/en/publication/66049"),
    ("R-07", "IEC 60079-0:2026 general requirements", "https://webstore.iec.ch/en/publication/71519"),
    ("R-08", "IEC 60079-1 flameproof enclosure d", "https://webstore.iec.ch/en/publication/621"),
    ("R-09", "IEC 60079-11:2023 intrinsic safety", "https://webstore.iec.ch/en/publication/60654"),
    ("R-10", "IEC 60079-17 inspection/maintenance", "https://webstore.iec.ch/en/publication/631"),
    ("R-11", "BSN SNI IEC 60079 catalog", "https://pesta.bsn.go.id/produk/index/471"),
    ("R-12", "NIST SP 800-82 Rev.3 OT Security", "https://csrc.nist.gov/pubs/sp/800/82/r3/final"),
    ("R-13", "OASIS MQTT Version 5.0", "https://docs.oasis-open.org/mqtt/mqtt/v5.0/mqtt-v5.0.html"),
    ("R-14", "TI BQ25185 datasheet", "https://www.ti.com/lit/ds/symlink/bq25185.pdf"),
    ("R-15", "TI TPS63020", "https://www.ti.com/product/TPS63020"),
    ("R-16", "TI LMR51450", "https://www.ti.com/product/LMR51450"),
    ("R-17", "Littelfuse miniSMDC resettable PTC datasheet", "https://m.littelfuse.com/~/media/electronics/datasheets/resettable_ptcs/littelfuse_ptc_minismdc_datasheet.pdf.pdf"),
    ("R-18", "Espressif ESP32-S3-WROOM-1/1U datasheet", "https://documentation.espressif.com/esp32-s3-wroom-1_wroom-1u_datasheet_en.pdf"),
    ("R-19", "EBYTE E22-900MM22S", "https://www.cdebyte.com/products/E22-900MM22S"),
    ("R-20", "IEC 62133-2 battery safety", "https://webstore.iec.ch/en/publication/32662"),
    ("R-21", "UN 38.3 test summary resources", "https://unece.org/transport/dangerous-goods/rev8-files"),
]
for sid, name, url in official_sources:
    para = doc.add_paragraph(style="Small Note")
    r = para.add_run(f"{sid}  {name}: ")
    set_run_font(r, size=9, bold=True, color=DARK)
    add_hyperlink(para, "Open official source", url)
    set_lang(para)

heading("15.4 MQ reference datasheets", 2)
mq_sources = [
    ("MQ-2", "https://www.winsen-sensor.com/d/files/manual/mq-2.pdf?v=1.0.24"),
    ("MQ-3B", "https://www.winsen-sensor.com/d/files/manual/mq-3b.pdf?searchid=6821"),
    ("MQ-4", "https://www.winsen-sensor.com/d/files/semiconductor/mq-4.pdf"),
    ("MQ-5", "https://www.winsen-sensor.com/d/files/manual/mq-5.pdf"),
    ("MQ-6", "https://www.winsen-sensor.com/d/files/manual/mq-6.pdf"),
    ("MQ-7B", "https://www.winsen-sensor.com/d/files/manual/mq-7b.pdf"),
    ("MQ-8", "https://www.winsen-sensor.com/d/files/semiconductor/mq-8.pdf"),
    ("MQ135", "https://www.winsen-sensor.com/d/files/mq135-%28ver1_6%29---manual.pdf"),
]
for name, url in mq_sources:
    para = doc.add_paragraph(style="Small Note")
    r = para.add_run(f"Winsen {name}: ")
    set_run_font(r, size=9, bold=True, color=DARK)
    add_hyperlink(para, "Open vendor datasheet", url)
    set_lang(para)
p("Semua MQ source di atas hanya reference candidate. Manufacturer/MPN/variant/lot yang benar-benar terpasang belum ditemukan pada EasyEDA/BOM/repo.", style="Small Note")

heading("15.5 Pernyataan akhir audit", 2)
callout("TRACEABILITY", "Setiap klaim final harus menunjuk salah satu: controlled source, controlled drawing/BOM, build artifact, measurement record, network/broker log, deployment revision, certificate database, atau signed FAT/SAT. Ketika evidence itu belum ada, status tetap TBD/NO-PROOF—bukan diasumsikan lulus.")
p("— Akhir laporan —", style="Small Note", align=WD_ALIGN_PARAGRAPH.CENTER)


# Document-level XML settings
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

# Ensure default document language and compatibility metadata.
styles_root = doc.styles.element
doc_defaults = styles_root.find(qn("w:docDefaults"))
if doc_defaults is not None:
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is not None:
        rpr = rpr_default.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            rpr_default.append(rpr)
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:val"), "id-ID")
        lang.set(qn("w:eastAsia"), "id-ID")

# Every table, including single-cell callouts, exposes an explicit first-row
# header to assistive technology. Data tables also repeat that row on split.
for _table in doc.tables:
    if _table.rows:
        set_repeat_table_header(_table.rows[0])
        set_row_cant_split(_table.rows[0])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
